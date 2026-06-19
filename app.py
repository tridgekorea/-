import streamlit as st
import pandas as pd
from docx import Document
import io
import requests
import json

# 1. 트릿지 테마 스타일 UI 설정
st.set_page_config(page_title="Tridge Data Solution Playbook", layout="wide")
st.title("📊 Tridge Data Solutions — 리포트 자동화 시스템")
st.markdown("### **\"절대값을 버리고, 시장 대비 위치를 측정하라\"**")
st.caption("샘표식품(주) 데이터 기반 구매 단가 진단 및 절감·방어 기회 분석 툴")

# 2. 사이드바 설정 (설정 및 API 입력창)
st.sidebar.header("⚙️ 분석 설정 및 API 연동")

# [중요] 발급받으신 클로드 API Key를 여기에 넣거나 화면에서 직접 입력합니다.
anthropic_key = st.sidebar.text_input("Anthropic (Claude) API Key를 입력하세요", type="password")

uploaded_file = st.sidebar.file_uploader("Tridge Explorer CSV 파일 업로드", type=['csv'])
target_company = st.sidebar.text_input("분석 대상 기업명 (예: 영인)", value="YOUNG IN")
cutoff_date = st.sidebar.date_input("이벤트 컷오프(기준) 날짜")

# 샘플 품목 검증용 필터링 단어
sample_product = st.sidebar.text_input("Mix Integrity 검증용 세부 제품명", value="DIVELLA")

if uploaded_file and target_company:
    # 3. 데이터 읽기 및 미리보기
    df = pd.read_csv(uploaded_file)
    st.subheader("📁 업로드된 데이터 미리보기 (상위 5개 행)")
    st.dataframe(df.head())

    try:
        # 4. 실제 트릿지 데이터 컬럼명 매핑 및 전처리
        df['Date'] = pd.to_datetime(df['Date'])
        cutoff = pd.to_datetime(cutoff_date)
        
        # Before / After 데이터 분할
        df_before = df[df['Date'] < cutoff]
        df_after = df[df['Date'] >= cutoff]

        # 물량가중 평균 단가(VWAP) 계산 함수: 금액(Value) / 물량(Volume)
        def get_vwap(data):
            if data['Volume'].sum() == 0: return 0
            return data['Value'].sum() / data['Volume'].sum()

        # 대상 기업 vs 시장 벤치마크(대상 제외) 분리
        # Importer 컬럼 또는 Raw Importer Name 컬럼 사용
        target_before = df_before[df_before['Importer'].str.contains(target_company, na=False, case=False)]
        target_after = df_after[df_after['Importer'].str.contains(target_company, na=False, case=False)]
        
        market_before = df_before[~df_before['Importer'].str.contains(target_company, na=False, case=False)]
        market_after = df_after[~df_after['Importer'].str.contains(target_company, na=False, case=False)]

        # 가중 단가 산출
        vw_target_before = get_vwap(target_before)
        vw_target_after = get_vwap(target_after)
        vw_market_before = get_vwap(market_before)
        vw_market_after = get_vwap(market_after)

        # 격차 및 성과 지표 계산
        gap_before = (vw_target_before / vw_market_before) - 1 if vw_market_before else 0
        gap_after = (vw_target_after / vw_market_after) - 1 if vw_market_after else 0
        gap_improvement = (gap_before - gap_after) * 100 # %p 개선량

        # 절감액 산출
        total_after_volume = target_after['Volume'].sum()
        saved_usd = total_after_volume * (vw_market_after - vw_target_after)

        # Mix Integrity 단일 품목 검증 (예: DIVELLA 건면 파스타 등)
        mix_target_before = target_before[target_before['Reported Product Name'].str.contains(sample_product, na=False, case=False)]
        mix_target_after = target_after[target_after['Reported Product Name'].str.contains(sample_product, na=False, case=False)]
        
        mix_before_price = get_vwap(mix_target_before)
        mix_after_price = get_vwap(mix_target_after)
        mix_change_rate = ((mix_after_price / mix_before_price) - 1) * 100 if mix_before_price else 0

        # 5. 화면 대시보드 UI 꾸미기
        st.divider()
        st.subheader(f"💡 {target_company} 성과 분석 결과 요약")
        
        # 큰 메트릭 UI로 이사님 보고용 지표 노출
        col1, col2, col3 = st.columns(3)
        col1.metric("시장 대비 단가 격차 개선량", f"{gap_improvement:+.1f}%p", help="양수일수록 시장 평균 대비 싸게 사기 시작했다는 뜻입니다.")
        col2.metric("시장 평균가 대비 총 절감액", f"${saved_usd:,.0f}", help="After 기간의 수입 물량에 개선된 단가차이를 곱한 금액입니다.")
        col3.metric("검증 품목 단가 변동률", f"{mix_change_rate:+.1f}%", help="Mix 무결성 확인을 위한 특정 단일 품목의 Before/After 단가 변동률입니다.")

        # 상세 데이터 테이블 제공
        st.markdown("#### **상세 지표 테이블**")
        summary_table = pd.DataFrame({
            "구분": ["대상 기업 단가 ($/kg)", "시장 벤치마크 단가 ($/kg)", "시장 대비 격차 (%)"],
            "Before (컷오프 전)": [f"${vw_target_before:.3f}", f"${vw_market_before:.3f}", f"{gap_before*100:+.1f}%"],
            "After (컷오프 후)": [f"${vw_target_after:.3f}", f"${vw_market_after:.3f}", f"{gap_after*100:+.1f}%"]
        })
        st.table(summary_table)

        # 6. 진짜 클로드 API 연동하여 서술형 분석 글 받아오기
        st.divider()
        st.subheader("🤖 Claude AI 전략 컨설팅 리포트 서술")
        
        ai_insight_text = "API Key를 입력하시면 클로드 AI가 자동으로 정밀 분석 단락을 서술합니다."
        
        if anthropic_key:
            with st.spinner("클로드 AI가 실거래 데이터를 기반으로 전략 보고서를 작성하고 있습니다..."):
                # Anthropic API 호출 (Claude 3.5 Sonnet 모델 사용)
                headers = {
                    "x-api-key": anthropic_key,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json"
                }
                
                prompt_content = f"""
                너는 글로벌 농식품 무역 데이터 분석 및 B2B 전략 컨설팅 전문가야. 
                파이썬이 계산해 준 아래의 실제 팩트 데이터를 바탕으로, 샘표식품 초기미팅 보고서 양식처럼 최고 경영진에게 보고할 수준의 전문적이고 논리적인 분석 리포트(인사이트 및 액션 아이템)를 한글로 작성해 줘.

                [분석 대상 및 데이터 개요]
                - 분석 대상 기업: {target_company}
                - 시장 대비 단가 격차 변화: {gap_improvement:.1f}%p 개선
                - 시장 평균가 대비 환산 절감액: ${saved_usd:,.0f}
                - Mix 무결성 검증: 특정 세부품목({sample_product})의 자체 단가가 {mix_change_rate:.1f}% 변동함.

                [작성 지시사항]
                1. 말투는 '~ 확인됩니다', '~ 판단됩니다'와 같이 냉철하고 격식 있는 컨설턴트 톤앤매너를 유지해라.
                2. 첫 단락은 '데이터 기반 진단 및 인사이트'로 작성하고, 이 단가 격차 개선이 우연이 아니며 믹스 무결성 검증을 통해 소싱 경쟁력이 강화되었음을 설명해라.
                3. 두 번째 단락은 '향후 원가 방어를 위한 Action Item'으로 작성하고, 트릿지 실거래 시세 데이터를 활용한 구체적인 공급망 관리 전략을 제안해라.
                """

                data = {
                    "model": "claude-3-5-sonnet-20241022",
                    "max_tokens": 1500,
                    "messages": [{"role": "user", "content": prompt_content}]
                }
                
                try:
                    response = requests.post("https://api.anthropic.com/v1/messages", headers=headers, json=data)
                    response_json = response.json()
                    ai_insight_text = response_json['content'][0]['text']
                    st.success("클로드 AI 분석 리포트 작성 완료!")
                except Exception as ai_err:
                    st.error(f"클로드 AI 연동 중 오류 발생: {ai_err}")
                    ai_insight_text = "API 통신 오류로 기본 문구로 대체됩니다."
                    
        st.write(ai_insight_text)

        # 7. 최종 고퀄리티 워드(.docx) 리포트 빌드 및 다운로드
        st.divider()
        st.subheader("📥 고퀄리티 전략 보고서 다운로드 (MS Word)")
        
        def create_word_report():
            doc = Document()
            # 타이틀 디자인 스타일 모방
            doc.add_heading('TRIDGE DATA SOLUTIONS', level=1)
            doc.add_heading(f'{target_company} 데이터 기반 구매 단가 진단 및 절감·방어 기회 분석', level=2)
            doc.add_paragraph('— 초기 미팅 전략 보고서 —\n')
            
            doc.add_heading('💡 핵심 요약 (Executive Summary)', level=3)
            p1 = doc.add_paragraph()
            p1.add_run(f"{target_company}의 수입 포트폴리오 데이터를 트릿지 실거래 데이터와 비교 분석했습니다. ").bold = True
            p1.add_run(f"비교의 정합성을 위해 대상 기업을 제외한 시장 벤치마크(ex-subject)를 구축하여 통제했습니다. 그 결과 이벤트 시점({cutoff_date}) 이후 시장 대비 단가 격차가 ")
            p1.add_run(f"{gap_improvement:.1f}%p 개선").bold = True
            p1.add_run(f"된 것이 확인되며, 이를 통한 시장 평균가 대비 총 절감 여력은 ")
            p1.add_run(f"${saved_usd:,.0f}").bold = True
            p1.add_run("로 추산됩니다.\n")
            
            doc.add_heading('📊 상세 분석 지표 (물량가중 평균 단가)', level=3)
            table = doc.add_table(rows=4, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '지표 구분'
            hdr_cells[1].text = 'Before (컷오프 전)'
            hdr_cells[2].text = 'After (컷오프 후)'
            
            row_cells1 = table.rows[1].cells
            row_cells1[0].text = f'{target_company} 단가'
            row_cells1[1].text = f'${vw_target_before:.2f}/kg'
            row_cells1[2].text = f'${vw_target_after:.2f}/kg'
            
            row_cells2 = table.rows[2].cells
            row_cells2[0].text = '시장 벤치마크 단가'
            row_cells2[1].text = f'${vw_market_before:.2f}/kg'
            row_cells2[2].text = f'${vw_market_after:.2f}/kg'
            
            row_cells3 = table.rows[3].cells
            row_cells3[0].text = '시장 대비 격차 (%)'
            row_cells3[1].text = f'{gap_before*100:+.1f}%'
            row_cells3[2].text = f'{gap_after*100:+.1f}%'

            doc.add_paragraph('\n')
            doc.add_heading('📝 세부 전략 분석 및 제안 (Claude AI Insight)', level=3)
            doc.add_paragraph(ai_insight_text)
            
            doc.add_paragraph('\n---\n본 보고서의 모든 수치는 트릿지가 확보한 글로벌 실거래 데이터에 근거하며, 전 거래 USD·KG 기준입니다.')
            
            bio = io.BytesIO()
            doc.save(bio)
            return bio.getvalue()

        docx_file = create_word_report()
        st.download_button(
            label="📄 완성된 전략 보고서 다운로드 (.docx)",
            data=docx_file,
            file_name=f"트릿지_{target_company}_전략보고서.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        st.error(f"데이터 연산 중 오류가 발생했습니다. 실제 CSV 파일 구조와 일치하는지 확인해 주세요. (에러 내용: {e})")
